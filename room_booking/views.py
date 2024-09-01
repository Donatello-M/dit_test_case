from django.db.models import Prefetch, Q
from django.http import HttpResponse
from docx import Document

from django.shortcuts import get_object_or_404
from drf_yasg.openapi import Parameter, IN_QUERY, FORMAT_DATETIME
from drf_yasg.utils import swagger_auto_schema
from rest_framework.response import Response
from rest_framework.views import APIView

from room_booking import mixins, models, serializers, utils


ROOM_MANUAL_PARAMETERS = [
    Parameter('start_date', IN_QUERY, 'Filter by date', type=FORMAT_DATETIME),
    Parameter('end_date', IN_QUERY, 'Filter by date', type=FORMAT_DATETIME)]


class RoomSchedule(mixins.AuthenticationMixin, APIView):
    """ REST API Расписание бронирования комнаты """

    @swagger_auto_schema(manual_parameters=ROOM_MANUAL_PARAMETERS)
    def get(self, request, room_name):
        room = get_object_or_404(models.Room, name=room_name)
        # По дефолту будет фильтровать бронь за текущую дату
        start_date, end_date = utils.get_filter_params(request)
        serializer = serializers.RoomSerializer(room, context={'start_date': start_date, 'end_date': end_date})
        return Response(serializer.data)


class RoomBooking(mixins.AuthenticationMixin, APIView):
    """ REST API Бронирование комнаты """

    @swagger_auto_schema(request_body=serializers.CreateReserveSerializer,
                         responses={201: serializers.CreateReserveSerializer(many=False)})
    def post(self, request):
        serializer = serializers.CreateReserveSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        request_data = serializer.validated_data
        room = request_data.get('room')
        serializer.save(reserved_by=request.user, room=room)
        return Response(serializer.data, status=201)


class BookingReportRetieve(mixins.AuthenticationMixin, APIView):
    """ Получение отчета по конкретной комнате """

    @swagger_auto_schema(manual_parameters=ROOM_MANUAL_PARAMETERS)
    def get(self, request, room_name):
        room = get_object_or_404(models.Room, name=room_name)
        start_date, end_date = utils.get_filter_params(request)
        reserves = models.Reserve.objects.filter(
            Q(room=room) & Q(start_time__gte=start_date) & Q(end_time__lte=end_date)).select_related(
            'reserved_by').iterator()
        doc = Document()
        doc.add_heading('Отчет', 0)
        doc.add_heading(room_name, level=1)

        for reserve in reserves:
            doc.add_heading(reserve.reserved_by, level=2)
            doc.add_heading(reserve.start_time, level=2)
            doc.add_heading(reserve.end_time, level=2)
            doc.add_paragraph(reserve.description)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="report_{room_name}.docx"'

        doc.save(response)

        return response


class BookingReportList(mixins.AuthenticationMixin, APIView):
    """ Получение отчета по всем комнатам """

    @swagger_auto_schema(manual_parameters=ROOM_MANUAL_PARAMETERS)
    def get(self, request):
        start_date, end_date = utils.get_filter_params(request)
        rooms = models.Room.objects.all().prefetch_related(
            Prefetch('room_reserves', queryset=models.Reserve.objects.filter(
                Q(start_time__gte=start_date) & Q(end_time__lte=end_date)).select_related(
                'reserved_by')))
        doc = Document()
        doc.add_heading('Отчет', 0)
        for room in rooms:
            doc.add_heading(room.name, level=1)

            for reserve in room.room_reserves.all():
                doc.add_heading(f'reserved_by: {reserve.reserved_by}', level=2)
                doc.add_heading(f'{reserve.start_time} - {reserve.end_time}', level=2)
                doc.add_paragraph(f'desciption: {reserve.description}')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="report.docx"'

        doc.save(response)

        return response
